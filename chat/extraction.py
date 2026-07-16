"""Text extraction from uploaded documents.

The agent's tools read plain text (`Document.content`), so a file uploaded in the
admin (a CV as PDF, a Word cover letter) is converted to text once, at upload time —
not on every chat turn. Extraction is best-effort: the admin can hand-edit the result
afterwards. Anything unreadable raises ValueError with a human-readable message, which
the admin form surfaces as a normal validation error.
"""

from io import BytesIO
from pathlib import Path

from docx import Document as DocxDocument
from pypdf import PdfReader

# Upload guardrails for the admin form. The size cap stays well under Django's 64MB
# body limit (raised for the PostHog proxy) — these are CV-sized documents, and the
# bytes are stored in the free-tier database, so small matters.
MAX_UPLOAD_BYTES = 10 * 1024 * 1024  # 10 MB
EXTENSION_CONTENT_TYPES = {
    ".pdf": "application/pdf",
    ".docx": "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
    ".txt": "text/plain",
    ".md": "text/markdown",
}


def content_type_for(filename: str) -> str:
    """The content type stored with the blob — picks the admin preview (PDFs render
    in an iframe; anything else is a download link)."""
    return EXTENSION_CONTENT_TYPES.get(Path(filename).suffix.lower(), "application/octet-stream")


def extract_text(data: bytes, filename: str) -> str:
    """Plain text from an uploaded file, chosen by extension.

    Raises ValueError when the format is unsupported, the file can't be parsed, or it
    yields no text at all (e.g. a scanned/image-only PDF — there's no OCR here)."""
    suffix = Path(filename).suffix.lower()
    if suffix not in EXTENSION_CONTENT_TYPES:
        supported = ", ".join(sorted(EXTENSION_CONTENT_TYPES))
        raise ValueError(f"Unsupported file type '{suffix or filename}' — use one of: {supported}")
    try:
        if suffix == ".pdf":
            reader = PdfReader(BytesIO(data))
            text = "\n".join(page.extract_text() or "" for page in reader.pages)
        elif suffix == ".docx":
            docx = DocxDocument(BytesIO(data))
            parts = [paragraph.text for paragraph in docx.paragraphs]
            # CVs in Word are often laid out as tables, which paragraphs don't cover.
            for table in docx.tables:
                parts.extend(" | ".join(cell.text for cell in row.cells) for row in table.rows)
            text = "\n".join(parts)
        else:  # .txt / .md
            text = data.decode("utf-8", errors="replace")
    except Exception as exc:  # noqa: BLE001 — any parser failure becomes a form error
        raise ValueError(f"Could not read '{filename}' — is the file valid?") from exc
    text = text.strip()
    if not text:
        raise ValueError(
            f"No text could be extracted from '{filename}' (a scanned/image-only PDF?). "
            "Paste the content in manually instead."
        )
    return text
